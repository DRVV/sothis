import streamlit as st
import io
from docx import Document

# Initialize session state for messages
if "messages" not in st.session_state:
    st.session_state.messages = []

# Sidebar for file upload
st.sidebar.title("Upload Options")
uploaded_file = st.sidebar.file_uploader("Upload a .docx file", type=["docx"])

if uploaded_file:
    # In reality, parse/transform content from the docx
    # We'll just store a dummy placeholder in session state
    st.session_state.doc_text = "Dummy content extracted from the uploaded docx..."
    st.sidebar.success("Docx file uploaded!")
else:
    # If no file is uploaded yet, ensure the key doesn't exist or is None
    st.session_state.doc_text = None

st.title("ChatGPT-like Local LLM Mock App")

# Display conversation so far
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# Chat input
user_input = st.chat_input("Type your message here...")
#if user_input:
if uploaded_file:
    # 1. Add user message to conversation
    user_input = "文字起こしをアップロードしました。議事録を作成してください。"
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.write(user_input)

    # 2. Create an LLM-like response
    if st.session_state.doc_text:
        # Example: incorporate doc content + user message
        assistant_reply = (
            f"Mock LLM answer that references uploaded doc content:\n\n"
            f"'{st.session_state.doc_text}'\n\n"
            f"User asked: '{user_input}'"
        )
    else:
        assistant_reply = f"Mock LLM answer with NO doc content. User asked: '{user_input}'"

    # 3. Add assistant response to session state
    st.session_state.messages.append({"role": "assistant", "content": assistant_reply})

    # 4. Display assistant response
    with st.chat_message("assistant"):
        st.write(assistant_reply)

        # If a file was uploaded, provide a download link for a docx processed by the LLM
        if st.session_state.doc_text:
            # Generate a .docx in memory with python-docx
            doc = Document()
            doc.add_paragraph("Below is the LLM's processed output:")
            doc.add_paragraph(assistant_reply)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_data = buffer.getvalue()

            st.download_button(
                label="Download Processed Docx",
                data=docx_data,
                file_name="processed_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Optionally, a button to download the *entire conversation* as Markdown
# if st.session_state.messages:
#     if st.button("Download conversation as Markdown"):
#         full_conv_md = ""
#         for msg in st.session_state.messages:
#             role = msg["role"].capitalize()
#             content = msg["content"]
#             full_conv_md += f"**{role}:** {content}\n\n"
        
#         st.download_button(
#             label="Download .md",
#             data=full_conv_md,
#             file_name="conversation.md",
#             mime="text/markdown"
#         )
