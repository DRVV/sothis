import streamlit as st
import io
from docx import Document

# Initialize conversation storage
if "messages" not in st.session_state:
    st.session_state.messages = []

# Track whether we've automatically processed the current file
# If user uploads multiple files, you could store each filename in session to handle them uniquely.
if "file_processed" not in st.session_state:
    st.session_state.file_processed = False

# Sidebar for file upload
st.sidebar.title("Upload Options")
uploaded_file = st.sidebar.file_uploader("Upload a .docx file", type=["docx"])

# If a file is uploaded, extract doc text (mock), and reset "file_processed" for each new file
if uploaded_file:
    # Real-world logic would parse the actual docx here
    st.session_state.doc_text = "Dummy content extracted from the uploaded docx..."
    # If this is a new file, mark as unprocessed so we trigger the automatic step
    if not st.session_state.file_processed:
        # We'll process it automatically below
        pass
else:
    st.session_state.doc_text = None
    st.session_state.file_processed = False

st.title("Auto-Processing File + ChatGPT-like UI")

# --- Automatic processing step (once per new file) ---
if st.session_state.doc_text and not st.session_state.file_processed:
    # 1) Create a system message
    system_msg = "I got a file, so I start working..."
    st.session_state.messages.append({"role": "system", "content": system_msg})
    
    # 2) "Process" the doc with a mock LLM
    assistant_reply = f"Here is a mock LLM result for the doc:\n\n{st.session_state.doc_text}"
    st.session_state.messages.append({"role": "assistant", "content": assistant_reply})
    
    st.session_state.file_processed = True  # Mark as done

# --- Display the conversation so far ---
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
        # Optionally, if this is from the assistant and a file was uploaded, show a download button
        # (depends on your desired UI; here's an example if you want to allow immediate download)
        if msg["role"] == "assistant" and st.session_state.doc_text:
            doc = Document()
            doc.add_paragraph("Below is the automatically processed output:")
            doc.add_paragraph(msg["content"])
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_data = buffer.getvalue()

            st.download_button(
                label="Download Processed Docx",
                data=docx_data,
                file_name="processed_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- Chat Input for further questions (optional) ---
user_input = st.chat_input("Ask a question or continue the conversation...")
if user_input:
    # Add user message
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.write(user_input)

    # Mock LLM reply (could incorporate doc text if desired)
    if st.session_state.doc_text:
        assistant_reply = (
            f"Mock LLM answer referencing the doc:\n{st.session_state.doc_text}\n\n"
            f"You asked: {user_input}"
        )
    else:
        assistant_reply = f"Mock LLM answer with no doc reference. You asked: {user_input}"

    st.session_state.messages.append({"role": "assistant", "content": assistant_reply})
    with st.chat_message("assistant"):
        st.write(assistant_reply)

# # --- Download entire conversation as Markdown (optional) ---
# if st.session_state.messages:
#     if st.button("Download conversation as Markdown"):
#         full_conv_md = ""
#         for msg in st.session_state.messages:
#             role = msg["role"].capitalize()
#             content = msg["content"]
#             full_conv_md += f"**{role}:** {content}\n\n"

#         st.download_button(
#             label="Download .md",
#             data=full_conv_md,
#             file_name="conversation.md",
#             mime="text/markdown"
#         )
